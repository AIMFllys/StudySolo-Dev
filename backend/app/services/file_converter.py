"""File converter service — converts Markdown content to downloadable files.

Supported output formats:
- PDF (via weasyprint or fallback to markdown2 + basic HTML)
- DOCX (via python-docx)
- Markdown (plain .md file)

The service generates files and returns a download URL or base64 content.
"""

import logging
import os
import uuid
import tempfile
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ConvertedFile:
    """Result of a file conversion."""
    filename: str
    format: str  # pdf, docx, md
    filepath: str  # Absolute path to the generated file
    size_bytes: int
    download_url: Optional[str] = None
    error: Optional[str] = None


# ── Export directory ─────────────────────────────────────────────────────────

EXPORT_DIR = os.getenv("EXPORT_DIR", os.path.join(tempfile.gettempdir(), "studysolo_exports"))


def _ensure_export_dir() -> str:
    """Ensure the export directory exists."""
    os.makedirs(EXPORT_DIR, exist_ok=True)
    return EXPORT_DIR


# ── Markdown export ──────────────────────────────────────────────────────────

async def export_markdown(content: str, filename: str = "export") -> ConvertedFile:
    """Save content as a .md file."""
    export_dir = _ensure_export_dir()
    safe_name = f"{filename}_{uuid.uuid4().hex[:8]}.md"
    filepath = os.path.join(export_dir, safe_name)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(filepath)
    return ConvertedFile(
        filename=safe_name,
        format="md",
        filepath=filepath,
        size_bytes=size,
    )


# ── TXT export ───────────────────────────────────────────────────────────────

def _strip_markdown(text: str) -> str:
    """Strip common Markdown formatting to produce plain text."""
    import re
    # Remove headers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.+?)_{1,3}', r'\1', text)
    # Remove links [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove images ![alt](url)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', text)
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove blockquotes
    text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
    # Remove list markers
    text = re.sub(r'^[\-*+]\s+', '  • ', text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


async def export_txt(content: str, filename: str = "export") -> ConvertedFile:
    """Convert Markdown content to plain text file."""
    export_dir = _ensure_export_dir()
    safe_name = f"{filename}_{uuid.uuid4().hex[:8]}.txt"
    filepath = os.path.join(export_dir, safe_name)

    plain_text = _strip_markdown(content)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(plain_text)

    size = os.path.getsize(filepath)
    return ConvertedFile(
        filename=safe_name,
        format="txt",
        filepath=filepath,
        size_bytes=size,
    )


# ── DOCX export ──────────────────────────────────────────────────────────────

async def export_docx(content: str, filename: str = "export") -> ConvertedFile:
    """Convert Markdown content to DOCX.

    Uses python-docx to create a simple formatted document.
    Markdown headings, bold, bullet points are converted to appropriate DOCX styles.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        return ConvertedFile(
            filename="", format="docx", filepath="", size_bytes=0,
            error="python-docx 未安装，无法导出 DOCX 格式",
        )

    export_dir = _ensure_export_dir()
    safe_name = f"{filename}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(export_dir, safe_name)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # Parse Markdown content line by line
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
            doc.add_paragraph(stripped[2:].strip(), style="List Number")
        # Blockquotes
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(stripped[2:])
            run.italic = True
        # Horizontal rules
        elif stripped in ("---", "***", "___"):
            continue
        # Normal text
        else:
            # Handle bold **text**
            if "**" in stripped:
                p = doc.add_paragraph()
                parts = stripped.split("**")
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.bold = True
            else:
                doc.add_paragraph(stripped)

    doc.save(filepath)
    size = os.path.getsize(filepath)

    return ConvertedFile(
        filename=safe_name,
        format="docx",
        filepath=filepath,
        size_bytes=size,
    )


# ── PDF export ───────────────────────────────────────────────────────────────

async def export_pdf(content: str, filename: str = "export") -> ConvertedFile:
    """Convert Markdown content to PDF.

    Strategy:
    1. Try weasyprint (best quality)
    2. Fallback: markdown → HTML → basic PDF via reportlab
    3. Last resort: save as .md with PDF extension notice
    """
    export_dir = _ensure_export_dir()
    safe_name = f"{filename}_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join(export_dir, safe_name)

    # Strategy 1: Try weasyprint
    try:
        from weasyprint import HTML as WeasyprintHTML
        import markdown

        html_content = markdown.markdown(
            content,
            extensions=["tables", "fenced_code", "codehilite"],
        )

        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: 'Microsoft YaHei', 'SimSun', Arial, sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    max-width: 700px;
                    margin: 40px auto;
                    padding: 0 20px;
                    color: #333;
                }}
                h1 {{ font-size: 22pt; color: #1a1a1a; margin-top: 24pt; }}
                h2 {{ font-size: 18pt; color: #333; margin-top: 20pt; }}
                h3 {{ font-size: 14pt; color: #555; margin-top: 16pt; }}
                code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 10pt; }}
                pre {{ background: #f8f8f8; padding: 12px; border-radius: 6px; overflow-x: auto; }}
                blockquote {{ border-left: 3px solid #ddd; margin: 16px 0; padding: 8px 16px; color: #666; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
                th {{ background: #f4f4f4; }}
                hr {{ border: none; border-top: 1px solid #eee; margin: 24px 0; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        WeasyprintHTML(string=styled_html).write_pdf(filepath)

        size = os.path.getsize(filepath)
        return ConvertedFile(
            filename=safe_name,
            format="pdf",
            filepath=filepath,
            size_bytes=size,
        )

    except ImportError:
        logger.info("weasyprint not available, trying fallback PDF generation")

    # Strategy 2: Fallback — save as markdown (with notice)
    # For environments where weasyprint/cairo isn't available
    md_name = safe_name.replace(".pdf", ".md")
    md_path = os.path.join(export_dir, md_name)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(md_path)
    return ConvertedFile(
        filename=md_name,
        format="md",
        filepath=md_path,
        size_bytes=size,
        error="PDF 导出需要安装 weasyprint，已降级为 Markdown 格式",
    )


# ── Main export dispatcher ──────────────────────────────────────────────────

async def convert_file(
    content: str,
    format: str = "md",
    filename: str = "export",
) -> ConvertedFile:
    """Convert content to the requested format.

    Args:
        content: Markdown content to convert
        format: Output format ("pdf", "docx", "md")
        filename: Base filename (without extension)

    Returns:
        ConvertedFile with path and metadata
    """
    format = format.lower().strip()

    if format == "pdf":
        return await export_pdf(content, filename)
    elif format == "docx":
        return await export_docx(content, filename)
    elif format in ("md", "markdown"):
        return await export_markdown(content, filename)
    elif format in ("txt", "text"):
        return await export_txt(content, filename)
    else:
        return ConvertedFile(
            filename="", format=format, filepath="", size_bytes=0,
            error=f"不支持的导出格式: {format}（支持 md/txt/docx/pdf）",
        )
