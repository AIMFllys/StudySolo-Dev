"""File parser service — extracts text from uploaded documents.

Supports PDF, DOCX, Markdown, and plain text formats.
Each parser returns a list of sections with heading + content for downstream chunking.
"""

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """A section extracted from a document."""
    heading: str = ""
    content: str = ""
    page: int | None = None
    level: int = 0  # heading level (1-6)


@dataclass
class ParsedDocument:
    """Full parsed result of a document."""
    filename: str = ""
    file_type: str = ""
    sections: list[ParsedSection] = field(default_factory=list)
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


def parse_file(filename: str, content: bytes) -> ParsedDocument:
    """Parse a file based on its extension and return structured text."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
        "txt": _parse_text,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: .{ext}（支持 pdf, docx, md, txt）")

    doc = parser(filename, content)
    doc.filename = filename
    doc.file_type = ext

    # Build full_text from sections
    doc.full_text = "\n\n".join(
        (f"## {s.heading}\n{s.content}" if s.heading else s.content)
        for s in doc.sections
        if s.content.strip()
    )

    return doc


def _parse_pdf(filename: str, content: bytes) -> ParsedDocument:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("需要安装 pypdf: pip install pypdf")

    reader = PdfReader(io.BytesIO(content))
    sections: list[ParsedSection] = []

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append(ParsedSection(
                heading=f"第 {page_num} 页",
                content=text.strip(),
                page=page_num,
            ))

    return ParsedDocument(sections=sections, metadata={"pages": len(reader.pages)})


def _parse_docx(filename: str, content: bytes) -> ParsedDocument:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = Document(io.BytesIO(content))
    sections: list[ParsedSection] = []
    current_heading = ""
    current_content: list[str] = []
    current_level = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if it's a heading
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Save previous section
            if current_content:
                sections.append(ParsedSection(
                    heading=current_heading,
                    content="\n".join(current_content),
                    level=current_level,
                ))
                current_content = []

            current_heading = text
            try:
                current_level = int(para.style.name.replace("Heading ", "").strip())
            except (ValueError, AttributeError):
                current_level = 1
        else:
            current_content.append(text)

    # Save last section
    if current_content:
        sections.append(ParsedSection(
            heading=current_heading,
            content="\n".join(current_content),
            level=current_level,
        ))

    return ParsedDocument(sections=sections)


def _parse_markdown(filename: str, content: bytes) -> ParsedDocument:
    """Extract sections from Markdown by splitting on headings."""
    text = content.decode("utf-8", errors="replace")
    lines = text.split("\n")
    sections: list[ParsedSection] = []
    current_heading = ""
    current_content: list[str] = []
    current_level = 0

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#"):
            # Save previous section
            if current_content:
                sections.append(ParsedSection(
                    heading=current_heading,
                    content="\n".join(current_content),
                    level=current_level,
                ))
                current_content = []

            # Parse heading level
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            current_heading = stripped.lstrip("#").strip()
            current_level = level
        else:
            if stripped:
                current_content.append(line)

    # Save last section
    if current_content:
        sections.append(ParsedSection(
            heading=current_heading,
            content="\n".join(current_content),
            level=current_level,
        ))

    return ParsedDocument(sections=sections)


def _parse_text(filename: str, content: bytes) -> ParsedDocument:
    """Parse plain text as a single section."""
    text = content.decode("utf-8", errors="replace")
    sections = [ParsedSection(content=text.strip())]
    return ParsedDocument(sections=sections)
